import os
from docx import Document
import re

def clean_text(text):
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_filtered_paragraphs(doc_path, min_length=150):
    filtered = []
    doc = Document(doc_path)
    for p in doc.paragraphs:
        text = clean_text(p.text)
        if text:
            if text.lower().startswith("about") or len(text) >= min_length:
                filtered.append(p)
    return filtered

def extract_s1(doc_path, min_length=150):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    dateline_found = False
    
    for i in range(len(paragraphs)):
        text = clean_text(paragraphs[i].text)
        if text.lower().startswith("dateline:"):
            dateline_found = True
            for j in range(i + 1, len(paragraphs)):
                next_paragraph = clean_text(paragraphs[j].text)
                if len(next_paragraph) >= 100:
                    return next_paragraph, j
    
    if not dateline_found:
        for i in range(len(paragraphs)):
            text = clean_text(paragraphs[i].text)
            if len(text) >= min_length:
                return text, i
    
    return "", -1  # Return empty string and invalid index if no qualifying paragraph found

def extract_s3_section(filtered_paragraphs, min_length=150):
    s3_text = ""
    collecting = False

    for paragraph in filtered_paragraphs:
        text = clean_text(paragraph.text)
        if text.lower().startswith("about"):
            if collecting:
                s3_text += "\n"
            collecting = True
            s3_text += "\n" + text + "\n"
        elif collecting:
            s3_text += text + " "
            if len(text) >= min_length:
                collecting = False
    
    return s3_text.strip()

def extract_s1_s2_s3(doc_path, min_length=150):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    filtered_paragraphs = extract_filtered_paragraphs(doc_path, min_length)
    
    # Extract S1
    s1, s1_index = extract_s1(doc_path, min_length)
    
    # Extract S3
    s3 = extract_s3_section(filtered_paragraphs, min_length)
    s3_index = -1
    
    if s3:
        for i in range(s1_index + 1, len(paragraphs)):
            if clean_text(paragraphs[i].text).lower().startswith("about"):
                s3_index = i
                break
    
    # Extract S2
    s2 = ""
    if s1_index != -1:
        s2_texts = []
        if s3 and s3_index != -1:
            for paragraph in paragraphs[s1_index + 1:s3_index]:
                s2_texts.append(paragraph.text)
        else:
            last_min_length_index = -1
            for i in range(s1_index + 1, len(paragraphs)):
                if len(clean_text(paragraphs[i].text)) >= min_length:
                    last_min_length_index = i
            if last_min_length_index == -1:
                last_min_length_index = len(paragraphs) - 1
            for paragraph in paragraphs[s1_index + 1:last_min_length_index + 1]:
                s2_texts.append(paragraph.text)
        s2 = "\n".join(s2_texts)

    return s1, s2, s3

def process_folder(input_folder_path, output_folder_path, min_length=150):
    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)
    
    for filename in os.listdir(input_folder_path):
        if filename.endswith(".doc") or filename.endswith(".docx"):
            file_path = os.path.join(input_folder_path, filename)
            base_name = os.path.splitext(filename)[0]
            
            # Count paragraphs that meet the min_length condition
            doc = Document(file_path)
            long_paragraphs = [p.text for p in doc.paragraphs if len(clean_text(p.text)) >= min_length]

            if len(long_paragraphs) == 1:
                # Save the single paragraph as S2
                s2_file = os.path.join(output_folder_path, f"{base_name}_S2.docx")
                doc_s2 = Document()
                doc_s2.add_paragraph(long_paragraphs[0])
                doc_s2.save(s2_file)
            else:
                # Run existing S1, S2, and S3 extraction logic
                s1, s2, s3 = extract_s1_s2_s3(file_path, min_length)
                
                if s1:
                    s1_file = os.path.join(output_folder_path, f"{base_name}_S1.docx")
                    doc_s1 = Document()
                    doc_s1.add_paragraph(s1)
                    doc_s1.save(s1_file)
                    
                if s3 or s3 == "":
                    s3_file = os.path.join(output_folder_path, f"{base_name}_S3.docx")
                    doc_s3 = Document()
                    if s3:
                        doc_s3.add_paragraph(s3)
                    doc_s3.save(s3_file)

                if s2:
                    s2_file = os.path.join(output_folder_path, f"{base_name}_S2.docx")
                    doc_s2 = Document()
                    doc_s2.add_paragraph(s2)
                    doc_s2.save(s2_file)

# Example usage:
input_folder_path = "input_folder_path"
output_folder_path = "output-folder-path"
min_length = 150
process_folder(input_folder_path, output_folder_path, min_length)
